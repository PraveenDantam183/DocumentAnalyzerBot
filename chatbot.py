import streamlit as st
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain_community.chat_models import ChatOpenAI
from langchain.chains.summarize import load_summarize_chain
from langchain.docstore.document import Document
import os
import time

# Load API key from Streamlit secrets
OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]

st.title("Document Analysis Chatbot")

# Initialize session state for conversation history and file processing
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None
if "llm" not in st.session_state:
    st.session_state.llm = None
if "summary" not in st.session_state:
    st.session_state.summary = None
if "last_query" not in st.session_state:
    st.session_state.last_query = ""

# Main area for file upload with multiple file types
st.header("Upload Your Document")
uploaded_file = st.file_uploader("Upload a PDF, Word Doc, or Image", type=["pdf", "docx", "jpg", "jpeg", "png"], key="file_uploader")

# Process the uploaded file and generate summary
if uploaded_file is not None:
    with st.spinner("🔍 **Analyzing Document...**"):
        time.sleep(2)  # Simulate processing time for visibility
        # Extract text based on file type
        if uploaded_file.type == "application/pdf":
            pdf_reader = PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif uploaded_file.type in ["image/jpeg", "image/png"]:
            image = Image.open(uploaded_file)
            text = pytesseract.image_to_string(image)

    # Convert text to a Document object for summarization
    doc = Document(page_content=text)

    # Split text into chunks for Q&A
    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n", r"\."],
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)

    # Generate embeddings and create vector store
    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    st.session_state.vector_store = FAISS.from_texts(chunks, embeddings)

    # Initialize the LLM
    st.session_state.llm = ChatOpenAI(
        openai_api_key=OPENAI_API_KEY,
        model_name="gpt-3.5-turbo",
        temperature=0.2,
        max_tokens=1500
    )

    # Generate summary
    with st.spinner("🔍 **Generating Summary...**"):
        summarize_chain = load_summarize_chain(st.session_state.llm, chain_type="stuff")
        st.session_state.summary = summarize_chain.run([doc])
    st.subheader("Document Summary")
    st.write(st.session_state.summary)

# Search bar and submit logic for Enter or Send button
if st.session_state.vector_store and st.session_state.llm:
    st.subheader("Search the Document")
    search_col1, search_col2 = st.columns([9, 1])
    with search_col1:
        search_query = st.text_input("Enter your search query here:", placeholder="e.g., What are the main terms?", key="search_input", on_change=lambda: st.session_state.last_query)
    with search_col2:
        submit_button = st.button("Send", key="send_button", on_click=lambda: st.session_state.last_query)

    # Trigger Q&A on Enter or Send button
    if st.session_state.last_query != search_query and search_query.strip():
        with st.spinner("🤔 **Thinking...**"):
            time.sleep(1)  # Simulate processing time for visibility
            # Perform similarity search
            docs = st.session_state.vector_store.similarity_search(search_query, k=5)

            # Load Q&A chain
            qa_chain = load_qa_chain(st.session_state.llm, chain_type="stuff")
            response = qa_chain.run(input_documents=docs, question=search_query)

            # Store the question and answer in session state
            st.session_state.chat_history.append({"question": search_query, "answer": response})
            st.session_state.last_query = search_query

    # Display conversation history in a chat-like format
    st.subheader("Conversation History")
    chat_container = st.container()
    with chat_container:
        for i, entry in enumerate(st.session_state.chat_history):
            st.markdown(f"""
                <div style='background-color: #2a2a2a; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>
                    <strong style='color: #4CAF50;'>Q{i+1}:</strong> {entry['question']}
                </div>
                <div style='background-color: #3a3a3a; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>
                    <strong style='color: #4CAF50;'>A{i+1}:</strong> {entry['answer']}
                </div>
            """, unsafe_allow_html=True)

    # Display the latest response
    if st.session_state.chat_history:
        st.subheader("Latest Search Results")
        st.write(st.session_state.chat_history[-1]["answer"])

# Add instructions
st.sidebar.info("Upload a PDF, Word Doc, or Image to start analyzing. Use the search bar and 'Send' button (or press Enter) to ask questions about the content!")

# Custom CSS for better UI and chat-like styling
st.markdown("""
    <style>
    .stFileUploader > div {
        padding: 10px;
        border: 2px dashed #4CAF50;
        border-radius: 5px;
        text-align: center;
        background-color: #1a1a1a;
    }
    .stTextInput > div > input {
        width: 100%;
        padding: 10px;
        font-size: 16px;
        border-radius: 5px;
        border: 1px solid #ccc;
        background-color: #2a2a2a;
        color: #fff;
    }
    .stButton > button {
        width: 100%;
        padding: 10px;
        font-size: 16px;
        border-radius: 5px;
        background-color: #4CAF50;
        color: white;
    }
    .stSpinner > div {
        color: #4CAF50;
        font-size: 20px;
        font-weight: bold;
    }
    .chat-container {
        max-height: 300px;
        overflow-y: auto;
        padding: 10px;
        background-color: #1a1a1a;
        border-radius: 5px;
    }
    </style>
""", unsafe_allow_html=True)

# Install additional dependencies if not present
try:
    import docx
except ImportError:
    st.warning("docx package not found. Installing it now...")
    os.system("pip install python-docx")
    import docx
try:
    import pytesseract
except ImportError:
    st.warning("pytesseract package not found. Installing it now...")
    os.system("pip install pytesseract")
    import pytesseract
try:
    from PIL import Image
except ImportError:
    st.warning("PIL package not found. Installing it now...")
    os.system("pip install Pillow")
    from PIL import Image